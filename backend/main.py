# app.py
from fastapi import FastAPI, File, UploadFile, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from google.cloud import vision
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from deep_translator import GoogleTranslator
import io, os, uuid, tempfile, json, logging, cv2, numpy as np
from PIL import Image
import pytesseract

# ---------------- APP SETUP ----------------
app = FastAPI(title="Handwriting OCR API v10 (OpenCV + Tesseract + EasyOCR + Google)")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Setup logging
logging.basicConfig(
    filename="ocr_api.log",
    format="%(asctime)s - %(levelname)s - %(message)s",
    level=logging.INFO
)

# ---------------- GOOGLE VISION INIT ----------------
client = None
GOOGLE_CREDENTIALS_JSON = os.getenv("GOOGLE_APPLICATION_CREDENTIALS_JSON")

try:
    if GOOGLE_CREDENTIALS_JSON:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".json", mode="w") as tmp_file:
            json.dump(json.loads(GOOGLE_CREDENTIALS_JSON), tmp_file)
            tmp_file.flush()
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = tmp_file.name
        client = vision.ImageAnnotatorClient()
        logging.info("Google Vision client initialized from environment variable.")
    else:
        local_key = "backend/write2doc-862dc6095d89.json"
        if os.path.exists(local_key):
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = local_key
            client = vision.ImageAnnotatorClient()
            logging.info("Using local Google Vision credentials.")
        else:
            logging.warning("No Google Vision credentials found. Google OCR may be disabled.")
except Exception as e:
    logging.error(f"Error initializing Google Vision client: {e}")

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ---------------- UTILITIES (preprocessing + deskew + cleanup) ----------------
def deskew(image: np.ndarray) -> np.ndarray:
    """Estimate rotation and deskew image."""
    coords = np.column_stack(np.where(image > 0))
    if coords.size == 0:
        return image
    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle
    (h, w) = image.shape[:2]
    center = (w // 2, h // 2)
    M = cv2.getRotationMatrix2D(center, angle, 1.0)
    rotated = cv2.warpAffine(image, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
    return rotated

def preprocess_image_bytes(image_bytes: bytes) -> bytes:
    """Enhanced preprocessing:
       - decode, convert to gray
       - contrast/brightness (convertScaleAbs)
       - bilateral filter (denoise)
       - adaptive threshold (binarize)
       - deskew
       - median blur
    """
    np_img = np.frombuffer(image_bytes, np.uint8)
    img = cv2.imdecode(np_img, cv2.IMREAD_COLOR)
    if img is None:
        raise ValueError("Decoded image is None")
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    # improve contrast/brightness
    gray = cv2.convertScaleAbs(gray, alpha=1.4, beta=15)
    # denoise while preserving edges
    gray = cv2.bilateralFilter(gray, d=9, sigmaColor=75, sigmaSpace=75)
    # adaptive threshold
    binarized = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                      cv2.THRESH_BINARY, 31, 2)
    # deskew
    deskewed = deskew(binarized)
    # smooth small noise
    deskewed = cv2.medianBlur(deskewed, 3)
    _, buffer = cv2.imencode('.png', deskewed)
    return buffer.tobytes()

def cleanup_text(txt: str) -> str:
    """Basic cleanup: normalize spaces and stray characters"""
    if not txt:
        return txt
    # replace multiple whitespace with single and strip
    lines = [line.rstrip() for line in txt.splitlines()]
    cleaned = "\n".join([ " ".join(line.split()) for line in lines if line.strip() ])
    return cleaned.strip()

def translate_text(text: str, target_lang="en"):
    """Translate OCR text if requested"""
    try:
        if not text.strip():
            return text
        return GoogleTranslator(source='auto', target=target_lang).translate(text)
    except Exception as e:
        logging.error(f"Translation failed: {e}")
        return text

# ---------------- OCR ENGINES ----------------
def google_ocr_image(content: bytes) -> str:
    if not client:
        raise RuntimeError("Google Vision client not initialized.")
    image = vision.Image(content=content)
    image_context = {"language_hints": ["ar", "en", "ha", "fr"]}
    response = client.document_text_detection(image=image, image_context=image_context)
    if response.error.message:
        raise Exception(response.error.message)
    return response.full_text_annotation.text.strip()

def google_ocr_pdf(content: bytes) -> str:
    if not client:
        raise RuntimeError("Google Vision client not initialized.")
    input_doc = vision.InputConfig(content=content, mime_type="application/pdf")
    feature = vision.Feature(type_=vision.Feature.Type.DOCUMENT_TEXT_DETECTION)
    image_context = {"language_hints": ["ar", "en", "ha", "fr"]}
    request = vision.AnnotateFileRequest(features=[feature], input_config=input_doc, image_context=image_context)
    result = client.batch_annotate_files(requests=[request])
    extracted = ""
    for response in result.responses:
        for page in response.responses:
            extracted += page.full_text_annotation.text + "\n\n"
    return extracted.strip()

def tesseract_ocr_from_bytes(content: bytes, lang="eng+ara+fra") -> str:
    try:
        # If TESSERACT_CMD env set (e.g., in container), ensure pytesseract uses it
        tcmd = os.getenv("TESSERACT_CMD")
        if tcmd:
            pytesseract.pytesseract.tesseract_cmd = tcmd
        np_img = np.frombuffer(content, np.uint8)
        img = cv2.imdecode(np_img, cv2.IMREAD_GRAYSCALE)
        # Use PIL image for pytesseract
        pil_img = Image.fromarray(img)
        config = "--psm 6"  # Assume a block of text
        text = pytesseract.image_to_string(pil_img, lang=lang, config=config)
        return text.strip()
    except Exception as e:
        logging.error(f"Tesseract OCR failed: {e}")
        return ""

def easyocr_process(content: bytes) -> str:
    try:
        import easyocr
        np_img = np.frombuffer(content, np.uint8)
        img = cv2.imdecode(np_img, cv2.IMREAD_COLOR)
        reader = easyocr.Reader(["ar", "en", "fr"], gpu=False)  # disable GPU by default in hosted envs
        results = reader.readtext(img, detail=0)
        return "\n".join(results)
    except Exception as e:
        logging.error(f"EasyOCR failed: {e}")
        return ""

# ---------------- SMART OCR CHAIN ----------------
def smart_ocr_chain(image_bytes: bytes, preferred: str = "google") -> (str, str):
    """Try preferred OCR then fallback chain.
       Returns (text, engine_used)
    """
    # try Google (only for images, for pdf we use google_pdf directly)
    engines_tried = []
    # first try preferred
    if preferred == "google" and client:
        try:
            engines_tried.append("google")
            txt = google_ocr_image(image_bytes)
            if txt.strip():
                return cleanup_text(txt), "google"
        except Exception as e:
            logging.warning(f"Google OCR failed: {e}")

    # next try tesseract
    try:
        engines_tried.append("tesseract")
        txt = tesseract_ocr_from_bytes(image_bytes)
        if txt.strip():
            return cleanup_text(txt), "tesseract"
    except Exception as e:
        logging.warning(f"Tesseract failed: {e}")

    # finally try easyocr
    try:
        engines_tried.append("easyocr")
        txt = easyocr_process(image_bytes)
        if txt.strip():
            return cleanup_text(txt), "easyocr"
    except Exception as e:
        logging.warning(f"EasyOCR failed: {e}")

    logging.error(f"All OCR attempts failed. Tried: {engines_tried}")
    return "", "none"

# ---------------- MAIN OCR ROUTE ----------------
@app.post("/ocr/")
async def ocr_file(
    file: UploadFile = File(...),
    engine: str = Query("google", enum=["google", "tesseract", "easyocr"]),
    translate: bool = Query(False),
    target_lang: str = Query("en")
):
    """Perform OCR and return extracted text + downloadable links"""
    try:
        filename = file.filename.lower()
        content = await file.read()
        logging.info(f"📥 Received file: {filename}, engine={engine}, translate={translate}")

        if filename.endswith((".png", ".jpg", ".jpeg")):
            processed_bytes = preprocess_image_bytes(content)
            extracted_text, used_engine = smart_ocr_chain(processed_bytes, preferred=engine)
        elif filename.endswith(".pdf"):
            # for pdf we rely on Google batch_annotate_files if configured
            try:
                extracted_text = google_ocr_pdf(content)
                used_engine = "google"
            except Exception as e:
                logging.warning(f"Google PDF OCR failed: {e} — attempting Tesseract page-by-page fallback.")
                # Basic PDF -> images fallback (requires poppler in container to use pdf2image)
                try:
                    from pdf2image import convert_from_bytes
                    pages = convert_from_bytes(content)
                    texts = []
                    for p in pages:
                        bio = io.BytesIO()
                        p.save(bio, format="PNG")
                        page_bytes = bio.getvalue()
                        pb = preprocess_image_bytes(page_bytes)
                        t, u = smart_ocr_chain(pb, preferred=engine)
                        texts.append(t)
                    extracted_text = "\n\n".join(texts)
                    used_engine = "tesseract" if extracted_text else "none"
                except Exception as e2:
                    logging.error(f"PDF fallback failed: {e2}")
                    return JSONResponse({"error": "PDF OCR failed (no suitable engine available)."}, status_code=500)
        else:
            return JSONResponse({"error": "Unsupported file type"}, status_code=400)

        if not extracted_text.strip():
            return JSONResponse({"message": "❌ No readable text found."}, status_code=400)

        if translate:
            extracted_text = translate_text(extracted_text, target_lang)

        uid = str(uuid.uuid4())[:8]
        pdf_path = os.path.join(OUTPUT_DIR, f"ocr_{uid}.pdf")
        docx_path = os.path.join(OUTPUT_DIR, f"ocr_{uid}.docx")

        # Create DOCX
        doc = Document()
        doc.add_heading("Extracted Handwritten Text", level=1)
        doc.add_paragraph(extracted_text)
        doc.save(docx_path)

        # Create PDF
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=letter)
        c.setFont("Helvetica", 10)
        y = 750
        for line in extracted_text.split("\n"):
            if any("\u0600" <= ch <= "\u06FF" for ch in line):
                c.drawRightString(550, y, line)
            else:
                c.drawString(72, y, line)
            y -= 15
            if y < 50:
                c.showPage()
                y = 750
        c.save()
        with open(pdf_path, "wb") as f:
            f.write(pdf_buffer.getvalue())

        base_url = os.getenv("FLY_APP_URL", os.getenv("RENDER_EXTERNAL_URL", "http://127.0.0.1:8000"))

        return {
            "status": "success",
            "uid": uid,
            "engine_used": used_engine,
            "translated": translate,
            "target_lang": target_lang,
            "text_length": len(extracted_text),
            "preview": extracted_text[:400] + "..." if len(extracted_text) > 400 else extracted_text,
            "pdf_url": f"{base_url}/download/pdf/{uid}",
            "docx_url": f"{base_url}/download/docx/{uid}"
        }

    except Exception as e:
        logging.exception("OCR failed unexpectedly")
        return JSONResponse({"error": str(e)}, status_code=500)

# ---------------- DOWNLOAD ROUTES ----------------
@app.get("/download/pdf/{uid}")
async def download_pdf(uid: str):
    path = os.path.join(OUTPUT_DIR, f"ocr_{uid}.pdf")
    if not os.path.exists(path):
        return JSONResponse({"error": "File not found"}, status_code=404)
    return FileResponse(path, filename=f"ocr_{uid}.pdf", media_type="application/pdf")

@app.get("/download/docx/{uid}")
async def download_docx(uid: str):
    path = os.path.join(OUTPUT_DIR, f"ocr_{uid}.docx")
    if not os.path.exists(path):
        return JSONResponse({"error": "File not found"}, status_code=404)
    return FileResponse(path, filename=f"ocr_{uid}.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.get("/health")
def health():
    return {"status": "ok", "message": "OCR API v10 healthy & running"}

@app.get("/")
def home():
    return {"message": "🚀 OCR API v10 (OpenCV + Tesseract + EasyOCR + Google) is live!"}
